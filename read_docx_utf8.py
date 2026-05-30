import sys
from docx import Document

doc = Document(sys.argv[1])
with open(sys.argv[2], 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        f.write(p.text + '\n')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                f.write(cell.text + '\n')
print("Successfully extracted to:", sys.argv[2])
